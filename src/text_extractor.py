"""
Модуль извлечения текста из документов различных форматов:
- DOCX (python-docx)
- PDF (PyMuPDF)
- Изображения (PNG, JPG) через EasyOCR
"""

from pathlib import Path
from typing import Optional
from abc import ABC, abstractmethod


class TextExtractor(ABC):
    """Базовый класс для извлечения текста"""
    
    @abstractmethod
    def extract(self, file_path: str) -> str:
        """Извлечь текст из файла"""
        pass
    
    @staticmethod
    def get_extension(file_path: str) -> str:
        """Получить расширение файла в нижнем регистре"""
        return Path(file_path).suffix.lower()


class DocxExtractor(TextExtractor):
    """Извлечение текста из DOCX файлов"""
    
    def extract(self, file_path: str) -> str:
        from docx import Document
        
        try:
            doc = Document(file_path)
            texts = []
            
            # Извлекаем текст из параграфов
            for para in doc.paragraphs:
                if para.text.strip():
                    texts.append(para.text)
            
            # Извлекаем текст из таблиц (важно для отчётов!)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            texts.append(cell.text)
            
            return "\n".join(texts)
        except Exception as e:
            raise RuntimeError(f"Ошибка чтения DOCX файла {file_path}: {e}")


class PdfExtractor(TextExtractor):
    """Извлечение текста из PDF файлов"""
    
    def extract(self, file_path: str) -> str:
        import fitz  # PyMuPDF
        
        try:
            doc = fitz.open(file_path)
            texts = []
            for page in doc:
                text = page.get_text()
                if text.strip():
                    texts.append(text)
            doc.close()
            return "\n".join(texts)
        except Exception as e:
            raise RuntimeError(f"Ошибка чтения PDF файла {file_path}: {e}")


class ImageExtractor(TextExtractor):
    """Извлечение текста из изображений через OCR (EasyOCR)"""
    
    def __init__(self, languages: list = None):
        """
        Инициализация OCR экстрактора
        
        Args:
            languages: Список языков для распознавания. 
                       По умолчанию ['ru', 'en'] для русского и английского
        """
        self.languages = languages if languages else ['ru', 'en']
        self._reader = None
    
    @property
    def reader(self):
        """Ленивая инициализация OCR читателя"""
        if self._reader is None:
            import easyocr
            self._reader = easyocr.Reader(self.languages, gpu=False)
        return self._reader
    
    def extract(self, file_path: str) -> str:
        try:
            results = self.reader.readtext(file_path)
            texts = [result[1] for result in results]
            return "\n".join(texts)
        except Exception as e:
            raise RuntimeError(f"Ошибка OCR обработки файла {file_path}: {e}")


class DocumentLoader:
    """
    Универсальный загрузчик документов.
    Автоматически определяет тип файла и выбирает нужный экстрактор.
    """
    
    SUPPORTED_EXTENSIONS = {
        '.docx': DocxExtractor,
        '.pdf': PdfExtractor,
        '.png': ImageExtractor,
        '.jpg': ImageExtractor,
        '.jpeg': ImageExtractor,
    }
    
    def __init__(self):
        self._extractors = {}
    
    def _get_extractor(self, file_path: str) -> TextExtractor:
        """Получить экстрактор для файла по расширению"""
        ext = TextExtractor.get_extension(file_path)
        
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Неподдерживаемый формат файла: {ext}")
        
        if ext not in self._extractors:
            self._extractors[ext] = self.SUPPORTED_EXTENSIONS[ext]()
        
        return self._extractors[ext]
    
    def load(self, file_path: str) -> str:
        """
        Загрузить текст из файла
        
        Args:
            file_path: Путь к файлу
            
        Returns:
            Извлечённый текст
        """
        if not Path(file_path).exists():
            raise FileNotFoundError(f"Файл не найден: {file_path}")
        
        extractor = self._get_extractor(file_path)
        return extractor.extract(file_path)
    
    def load_multiple(self, file_paths: list[str], progress_callback=None) -> dict[str, str]:
        """
        Загрузить текст из нескольких файлов
        
        Args:
            file_paths: Список путей к файлам
            progress_callback: Callback для отображения прогресса (current, total, filename)
            
        Returns:
            Словарь {путь_к_файлу: текст}
        """
        results = {}
        total = len(file_paths)
        
        for i, file_path in enumerate(file_paths):
            try:
                text = self.load(file_path)
                results[file_path] = text
                
                if progress_callback:
                    progress_callback(i + 1, total, file_path)
                    
            except Exception as e:
                # Логируем ошибку, но продолжаем обработку остальных файлов
                print(f"Ошибка обработки файла {file_path}: {e}")
                results[file_path] = ""  # Пустой текст для ошибочных файлов
                
                if progress_callback:
                    progress_callback(i + 1, total, file_path)
        
        return results
