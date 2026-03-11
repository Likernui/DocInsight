#!/usr/bin/env python3
"""Скрипт для проверки извлечения текста из DOCX"""

import sys
from pathlib import Path
from docx import Document

if len(sys.argv) < 2:
    print("Использование: python check_docx.py <путь_к_файлу.docx>")
    sys.exit(1)

file_path = sys.argv[1]

if not Path(file_path).exists():
    print(f"Файл не найден: {file_path}")
    sys.exit(1)

print(f"Проверка файла: {file_path}")
print("=" * 60)

# Метод 1: python-docx
print("\n📄 Метод 1: python-docx")
doc = Document(file_path)
text_docx = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
print(f"Параграфов: {len(doc.paragraphs)}")
print(f"Символов: {len(text_docx)}")

# Метод 2: PyMuPDF
print("\n📄 Метод 2: PyMuPDF")
import fitz
doc_pdf = fitz.open(file_path)
text_pymupdf = ""
for page in doc_pdf:
    text_pymupdf += page.get_text()
doc_pdf.close()
print(f"Символов: {len(text_pymupdf)}")

# Метод 3: Прямое чтение XML
print("\n📄 Метод 3: XML (прямое чтение)")
import zipfile
import xml.etree.ElementTree as ET

texts = []
with zipfile.ZipFile(file_path, 'r') as zip_doc:
    print(f"Файлы в архиве: {zip_doc.namelist()[:10]}...")
    
    # Читаем document.xml
    if 'word/document.xml' in zip_doc.namelist():
        xml_content = zip_doc.read('word/document.xml')
        root = ET.fromstring(xml_content)
        
        # Ищем весь текст
        for elem in root.iter():
            if elem.text and elem.text.strip():
                texts.append(elem.text.strip())

xml_text = "\n".join(texts)
print(f"Символов: {len(xml_text)}")

# Выводим первые 1000 символов
print("\n" + "=" * 60)
print("ПЕРВЫЕ 1000 СИМВОЛОВ (любой метод):")
print("=" * 60)
print(text_docx[:1000])

# Статистика
print("\n" + "=" * 60)
print("СТАТИСТИКА:")
print("=" * 60)
print(f"python-docx: {len(text_docx)} символов")
print(f"PyMuPDF: {len(text_pymupdf)} символов")
print(f"XML: {len(xml_text)} символов")
